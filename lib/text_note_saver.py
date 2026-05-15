"""
长文本消息 → DOCX 落盘到 ARCHIVE_DIR。

典型场景:微信 "笔记" / 调研纪要 / 公众号正文 复制出来粘贴发送过来的长文本。
跟 wechat_mp_fetcher 输出格式对齐,下游 pipeline 统一处理 DOCX。

判定:
- 长度 >= NOTE_MIN_LENGTH(默认 200,可 .env 覆盖)的纯文本(无 URL)才触发落盘
- 短文本仍走"已识别为 文本消息"那条路径,避免每个 hi / 测试 都落盘污染目录
"""
from __future__ import annotations

import datetime as _dt
import os
import re
from pathlib import Path

from docx import Document

from lib.downloader import get_archive_dir


# 触发落盘的最低字符数;可被 .env 里 NOTE_MIN_LENGTH 覆盖
_DEFAULT_MIN_LENGTH = 200


def get_min_length() -> int:
    """读 .env 里的 NOTE_MIN_LENGTH,非法或缺失则用默认值"""
    val = os.getenv("NOTE_MIN_LENGTH", "").strip()
    if val.isdigit():
        return int(val)
    return _DEFAULT_MIN_LENGTH


def is_long_enough(content: str) -> bool:
    """判断这条文本够不够长,值得落盘"""
    return len(content.strip()) >= get_min_length()


def _sanitize_title(title: str, max_len: int = 60) -> str:
    """文件名安全化,去掉 Windows 非法字符 + 控制字符,空格转下划线"""
    if not title:
        return ""
    title = re.sub(r'[\\/*?:"<>|\n\r\t]', "", title)
    title = title.strip().replace(" ", "_")
    return title[:max_len]


def _extract_title(content: str, max_title_chars: int = 60) -> str:
    """
    从文本内容里抠"标题":
    - 第一行非空且 <= max_title_chars 的内容,直接用
    - 第一行超长则取它的前 30 字作摘要
    - 都没有返回空串
    """
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        if len(line) <= max_title_chars:
            return line
        return line[:30]
    return ""


def save_as_docx(content: str, body: dict) -> tuple[Path, int, str]:
    """
    长文本 → DOCX 落盘到 ARCHIVE_DIR。

    Args:
        content: 完整 text.content
        body: 完整 msgbody,用于抠 userid 作为元信息

    Returns:
        (落盘路径, 字节数, 抠出的标题)

    Raises:
        OSError: 写盘失败
    """
    title = _extract_title(content)
    archive_dir = get_archive_dir()
    archive_dir.mkdir(parents=True, exist_ok=True)

    ts = _dt.datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = _sanitize_title(title)
    filename = f"{ts}_{safe}.docx" if safe else f"{ts}.docx"
    path = archive_dir / filename

    # 极端边界:同一秒落两个文件
    if path.exists():
        for i in range(1, 100):
            cand = path.with_suffix(f".{i}{path.suffix}")
            if not cand.exists():
                path = cand
                break

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    # 元信息
    meta_lines = [
        "来源:微信文本消息(复制 / 长文本转发)",
        f"接收时间:{_dt.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ]
    userid = (body.get("from") or {}).get("userid")
    if userid:
        meta_lines.append(f"发送人 userid:{userid}")
    doc.add_paragraph("\n".join(meta_lines))
    doc.add_paragraph("=" * 40)

    # 正文按换行分段
    for paragraph in content.split("\n"):
        para = paragraph.strip()
        if para:
            doc.add_paragraph(para)

    doc.save(str(path))
    return path, path.stat().st_size, title
