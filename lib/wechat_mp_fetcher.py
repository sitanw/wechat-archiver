"""
公众号文章抓取 → DOCX

复用 AceCamp scraper 的核心套路:
- 异步 Playwright(chromium)
- 反检测 init_script
- python-docx 构造文档,图片 httpx 拉回嵌入

公众号特化:
- 标题 / 公众号名 / 发布时间 / 正文容器都是固定 ID(#activity-name / #js_name / #publish_time / #js_content)
- 图片可能在 src 也可能在 data-src(懒加载),先滚到底触发加载再抽取
- 不需要登录 / cookie / captcha
"""
from __future__ import annotations

import asyncio
import datetime as _dt
import os
import random
import re
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from docx.shared import Inches
from playwright.async_api import async_playwright, TimeoutError as PWTimeoutError

from lib.downloader import get_archive_dir


class LinkFetchError(Exception):
    """公众号抓取 / DOCX 构造失败的统一异常"""


def _include_images() -> bool:
    """读 .env 里的 INCLUDE_IMAGES,默认 False(公众号图多是广告/二维码,信噪比低)"""
    val = os.getenv("INCLUDE_IMAGES", "").strip().lower()
    return val in ("1", "true", "yes", "y", "on")


# 单次抓取总预算(从 launch 到 DOCX 落盘),超过算失败
TOTAL_TIMEOUT_SECONDS = 60

# 反检测 — 抄自 AceCamp scraper
_USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:125.0) Gecko/20100101 Firefox/125.0",
]

_STEALTH_INIT_SCRIPT = """
Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3] });
window.chrome = { runtime: {} };
Object.defineProperty(navigator, 'languages', { get: () => ['zh-CN', 'zh', 'en'] });
"""

# 公众号文章 DOM 提取脚本
# 返回 {title, author, date, blocks: [{type: 'text'|'image', value|url}]}
_EXTRACT_JS = r"""
() => {
    const result = { title: '', author: '', date: '', blocks: [] };

    // 标题
    const titleEl = document.querySelector('#activity-name, h1#activity-name, h2.rich_media_title, h1.rich_media_title');
    if (titleEl) result.title = (titleEl.innerText || '').trim();

    // 公众号名(用作 author)
    const authorEl = document.querySelector('#js_name, a#js_name, .rich_media_meta_nickname');
    if (authorEl) result.author = (authorEl.innerText || '').trim();

    // 发布时间
    const dateEl = document.querySelector('#publish_time, em#publish_time, em.rich_media_meta_text');
    if (dateEl) result.date = (dateEl.innerText || '').trim();

    // 正文容器:固定 ID,这是公众号文章页最稳定的锚点
    const content = document.querySelector('#js_content, .rich_media_content');
    if (!content) return result;

    // 遍历叶子节点,生成 text / image blocks
    const seen = new Set();
    const els = content.querySelectorAll('p, section, div, span, h1, h2, h3, h4, img, li, blockquote');

    for (const el of els) {
        if (el.tagName.toLowerCase() === 'img') {
            // 公众号图片优先 data-src(懒加载),fallback 到 src
            let src = el.getAttribute('data-src') || el.src || '';
            // 跳过 base64 缩略图 / 1x1 占位
            if (!src || src.startsWith('data:') || src.length < 20) continue;
            if (seen.has(src)) continue;
            seen.add(src);
            result.blocks.push({ type: 'image', url: src });
            continue;
        }

        // 只取叶子节点的文本,避免父子重复
        const hasBlockChild = el.querySelector('p, section, div, h1, h2, h3, h4, li, blockquote');
        if (hasBlockChild) continue;

        const text = (el.innerText || '').trim();
        if (!text || text.length < 2) continue;
        if (seen.has(text)) continue;
        seen.add(text);
        result.blocks.push({ type: 'text', value: text });
    }
    return result;
}
"""


def _sanitize_title(title: str, max_len: int = 60) -> str:
    """文件名安全化,去掉 Windows 非法字符 + 控制字符,空格转下划线"""
    if not title:
        return ""
    title = re.sub(r'[\\/*?:"<>|\n\r\t]', "", title)
    title = title.strip().replace(" ", "_")
    return title[:max_len]


def _build_filename(title: str) -> str:
    """
    {date}_{HHMMSS}_{title}.docx;无标题则只保留 {date}_{HHMMSS}。
    若 title 开头是日期(罕见,公众号文章标题通常不带),用它作 prefix,strip 出 title。
    """
    from lib.tag_parser import extract_leading_date

    embedded_date, title = extract_leading_date(title or "")
    date_str = embedded_date or _dt.datetime.now().strftime("%Y-%m-%d")
    time_str = _dt.datetime.now().strftime("%H%M%S")
    ts = f"{date_str}_{time_str}"
    safe = _sanitize_title(title)
    return f"{ts}_{safe}.docx" if safe else f"{ts}.docx"


async def _scroll_to_bottom(page, step_px: int = 800, max_steps: int = 30):
    """滚到底,触发图片懒加载。每步之间留点时间让 IntersectionObserver 工作。"""
    last_height = -1
    for _ in range(max_steps):
        height = await page.evaluate("document.body.scrollHeight")
        if height == last_height:
            break
        await page.evaluate(f"window.scrollBy(0, {step_px})")
        await asyncio.sleep(0.2)
        last_height = height
    # 滚回顶部,避免影响最终页面状态(无所谓但更整洁)
    await page.evaluate("window.scrollTo(0, 0)")


async def _build_docx(article: dict, target_path: Path) -> int:
    """根据 article 字典写 DOCX,返回字节数"""
    doc = Document()
    if article.get("title"):
        doc.add_heading(article["title"], level=1)

    meta_lines = []
    if article.get("author"):
        meta_lines.append(f"公众号:{article['author']}")
    if article.get("date"):
        meta_lines.append(f"发布时间:{article['date']}")
    if article.get("url"):
        meta_lines.append(f"原文:{article['url']}")
    if meta_lines:
        doc.add_paragraph("\n".join(meta_lines))
        doc.add_paragraph("=" * 40)

    include_imgs = _include_images()

    # 不带图:跳过所有 image block,直接写文字。省一个 httpx client 的开销
    if not include_imgs:
        for block in article.get("blocks", []):
            if block.get("type") == "text":
                val = block.get("value", "").strip()
                if val:
                    doc.add_paragraph(val)
        doc.save(str(target_path))
        return target_path.stat().st_size

    # 带图:正文文字 + 图片 httpx 拉回嵌入
    async with httpx.AsyncClient(timeout=15, follow_redirects=True) as client:
        for block in article.get("blocks", []):
            btype = block.get("type")
            if btype == "text":
                val = block.get("value", "").strip()
                if val:
                    doc.add_paragraph(val)
            elif btype == "image":
                img_url = block.get("url", "")
                if not img_url:
                    continue
                try:
                    resp = await client.get(img_url)
                    resp.raise_for_status()
                    img_stream = BytesIO(resp.content)
                    doc.add_picture(img_stream, width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[图片加载失败: {img_url} — {type(e).__name__}]")

    doc.save(str(target_path))
    return target_path.stat().st_size


async def fetch_and_save_as_docx(url: str, body: dict) -> tuple[Path, int, str]:
    """
    公众号 URL → DOCX 落盘到 ARCHIVE_DIR。

    Args:
        url: mp.weixin.qq.com/... 完整链接
        body: 完整 msgbody,目前未直接用,留参数对称

    Returns:
        (落盘路径, 字节数, 抓到的标题)

    Raises:
        LinkFetchError: 加载 / 抽取 / DOCX 构造任一步失败
    """
    async def _do():
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(
                headless=True,
                args=[
                    "--disable-blink-features=AutomationControlled",
                    "--no-sandbox",
                    "--disable-infobars",
                ],
            )
            try:
                ctx = await browser.new_context(
                    user_agent=random.choice(_USER_AGENTS),
                    viewport={"width": 1366, "height": 900},
                    locale="zh-CN",
                    timezone_id="Asia/Shanghai",
                )
                await ctx.add_init_script(_STEALTH_INIT_SCRIPT)
                page = await ctx.new_page()

                # 1. 加载页面
                try:
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                except PWTimeoutError as e:
                    raise LinkFetchError(f"页面加载超时(30s): {url}") from e

                # 2. 等正文容器出现(公众号文章页面 #js_content 是稳定锚点)
                try:
                    await page.wait_for_selector(
                        "#js_content, .rich_media_content", timeout=15000
                    )
                except PWTimeoutError as e:
                    raise LinkFetchError(
                        f"未找到正文容器 #js_content,可能不是公众号文章或被反爬: {url}"
                    ) from e

                # 3. 滚到底触发图片懒加载
                await _scroll_to_bottom(page)
                await asyncio.sleep(1)  # 让最后一批图加载

                # 4. 抽取
                article = await page.evaluate(_EXTRACT_JS)
            finally:
                await browser.close()

        article["url"] = url
        blocks = article.get("blocks", [])
        if not blocks:
            raise LinkFetchError(f"正文为空,可能页面结构变了: {url}")

        # 5. 构造 DOCX,落盘
        archive_dir = get_archive_dir()
        archive_dir.mkdir(parents=True, exist_ok=True)
        path = archive_dir / _build_filename(article.get("title", ""))
        # 同名冲突(罕见)处理
        if path.exists():
            for i in range(1, 100):
                cand = path.with_suffix(f".{i}{path.suffix}")
                if not cand.exists():
                    path = cand
                    break

        size = await _build_docx(article, path)
        return path, size, article.get("title", "")

    try:
        return await asyncio.wait_for(_do(), timeout=TOTAL_TIMEOUT_SECONDS)
    except asyncio.TimeoutError as e:
        raise LinkFetchError(
            f"总耗时超过 {TOTAL_TIMEOUT_SECONDS}s,放弃: {url}"
        ) from e
