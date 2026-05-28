"""
通用网页抓取 → DOCX(适用于 web_other 类 URL)。

架构跟 wechat_mp_fetcher 一致,只是不写 site-specific selectors,而是用
**trafilatura** 这个开源正文抽取库自动识别 title / author / date / body。

覆盖目标:中文财经新闻类站点 — 华尔街见闻 / 36氪 / 财联社 / 雪球 / 虎嗅 /
新浪财经 / Substack / 个人博客 等约 90%+ 标准文章页。

不适用:
- 公众号(走专用 wechat_mp_fetcher,质量更高)
- 有道云 / 腾讯文档 / 网盘(需要登录态,trafilatura 拿不到)
- 视频站(B 站 / 抖音):没有正文,主要是视频
- 付费墙(摘要能拿,完整正文拿不到)

实测发现某站质量不行 → 单独写 site-specific override(可参考 wechat_mp_fetcher
那套手写 selector 的模式)。
"""
from __future__ import annotations

import asyncio
import datetime as _dt
import os
import random
import re
from pathlib import Path

import trafilatura
from docx import Document
from playwright.async_api import async_playwright, TimeoutError as PWTimeoutError

from lib.downloader import get_archive_dir


class LinkFetchError(Exception):
    """通用网页抓取 / DOCX 构造失败的统一异常"""


# 单次抓取总预算(从 launch 到 DOCX 落盘),超过算失败
TOTAL_TIMEOUT_SECONDS = 60

# 反检测 — 跟 wechat_mp_fetcher 同款
_USER_AGENTS = [
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/123.0.0.0 Safari/537.36",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:125.0) Gecko/20100101 Firefox/125.0",
]

_STEALTH_INIT_SCRIPT = """
Object.defineProperty(navigator, 'webdriver', { get: () => undefined });
Object.defineProperty(navigator, 'plugins', { get: () => [1, 2, 3] });
window.chrome = { runtime: {} };
Object.defineProperty(navigator, 'languages', { get: () => ['zh-CN', 'zh', 'en'] });
"""


def _sanitize_title(title: str, max_len: int = 60) -> str:
    if not title:
        return ""
    title = re.sub(r'[\\/*?:"<>|\n\r\t]', "", title)
    title = title.strip().replace(" ", "_")
    return title[:max_len]


def _build_filename(title: str) -> str:
    """{date}_{HHMMSS}_{title}.docx — 跟 wechat_mp_fetcher 输出格式对齐"""
    from lib.tag_parser import extract_leading_date

    embedded_date, title = extract_leading_date(title or "")
    date_str = embedded_date or _dt.datetime.now().strftime("%Y-%m-%d")
    time_str = _dt.datetime.now().strftime("%H%M%S")
    ts = f"{date_str}_{time_str}"
    safe = _sanitize_title(title)
    return f"{ts}_{safe}.docx" if safe else f"{ts}.docx"


async def _scroll_to_bottom(page, step_px: int = 800, max_steps: int = 30):
    """滚到底,触发图片/正文懒加载"""
    last_height = -1
    for _ in range(max_steps):
        height = await page.evaluate("document.body.scrollHeight")
        if height == last_height:
            break
        await page.evaluate(f"window.scrollBy(0, {step_px})")
        await asyncio.sleep(0.2)
        last_height = height
    await page.evaluate("window.scrollTo(0, 0)")


def _extract_with_trafilatura(html: str, url: str) -> dict:
    """
    用 trafilatura 抽取正文 + 元信息。返回 dict:
        {title, author, date, sitename, text}
    抽不出来抛 LinkFetchError。
    """
    # 抽正文文本(去广告 / 评论 / 导航,保留段落结构)
    text = trafilatura.extract(
        html,
        url=url,
        include_comments=False,
        include_tables=True,        # 财经页面常有数据表
        include_formatting=False,
        favor_recall=True,          # 宁可多抓不漏
        output_format="txt",
    )
    if not text or len(text.strip()) < 50:
        raise LinkFetchError(
            f"trafilatura 抽不出正文(长度 {len(text or '')} < 50 字符),"
            f"可能是付费墙 / SPA 未渲染 / 站点结构特殊"
        )

    # 抽 metadata(title / author / date / sitename)
    meta_doc = trafilatura.extract_metadata(html, default_url=url)
    title = getattr(meta_doc, "title", "") if meta_doc else ""
    author = getattr(meta_doc, "author", "") if meta_doc else ""
    date = getattr(meta_doc, "date", "") if meta_doc else ""
    sitename = getattr(meta_doc, "sitename", "") if meta_doc else ""

    return {
        "title": (title or "").strip(),
        "author": (author or "").strip(),
        "date": (date or "").strip(),
        "sitename": (sitename or "").strip(),
        "text": text.strip(),
    }


async def _build_docx(article: dict, target_path: Path) -> int:
    """根据 article 字典写 DOCX,返回字节数"""
    doc = Document()
    if article.get("title"):
        doc.add_heading(article["title"], level=1)

    # 元信息块
    meta_lines = []
    if article.get("sitename"):
        meta_lines.append(f"站点:{article['sitename']}")
    if article.get("author"):
        meta_lines.append(f"作者:{article['author']}")
    if article.get("date"):
        meta_lines.append(f"发布时间:{article['date']}")
    if article.get("url"):
        meta_lines.append(f"原文:{article['url']}")
    if meta_lines:
        doc.add_paragraph("\n".join(meta_lines))
        doc.add_paragraph("=" * 40)

    # 正文按段落分(trafilatura 已经做了清洗,每段一行)
    for paragraph in article.get("text", "").split("\n"):
        para = paragraph.strip()
        if para:
            doc.add_paragraph(para)

    doc.save(str(target_path))
    return target_path.stat().st_size


async def fetch_and_save_as_docx(url: str, body: dict) -> tuple[Path, int, str]:
    """
    通用网页 URL → DOCX 落盘到 ARCHIVE_DIR。

    Returns:
        (path, size, title)

    Raises:
        LinkFetchError: 加载 / 抽取 / 落盘任一步失败
    """
    async def _do():
        async with async_playwright() as pw:
            browser = await pw.chromium.launch(
                headless=True,
                args=[
                    "--disable-blink-features=AutomationControlled",
                    "--no-sandbox",
                    "--disable-infobars",
                ],
            )
            try:
                ctx = await browser.new_context(
                    user_agent=random.choice(_USER_AGENTS),
                    viewport={"width": 1366, "height": 900},
                    locale="zh-CN",
                    timezone_id="Asia/Shanghai",
                )
                await ctx.add_init_script(_STEALTH_INIT_SCRIPT)
                page = await ctx.new_page()

                # 1. 加载
                try:
                    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                except PWTimeoutError as e:
                    raise LinkFetchError(f"页面加载超时(30s): {url}") from e

                # 2. 等 SPA 渲染 + 滚到底触发懒加载
                # 通用站点没有"等某个 selector 出现"的标准 anchor,等 networkidle 比较稳
                try:
                    await page.wait_for_load_state("networkidle", timeout=10000)
                except PWTimeoutError:
                    pass  # 部分站永远 networkidle 不到(WebSocket / 长轮询),不致命

                await _scroll_to_bottom(page)
                await asyncio.sleep(1)

                # 3. 拿完整 HTML
                html = await page.content()
            finally:
                await browser.close()

        # 4. trafilatura 抽取
        article = _extract_with_trafilatura(html, url)
        article["url"] = url

        # 5. 构造 DOCX 落盘
        archive_dir = get_archive_dir()
        archive_dir.mkdir(parents=True, exist_ok=True)
        path = archive_dir / _build_filename(article.get("title", ""))
        # 同名冲突
        if path.exists():
            for i in range(1, 100):
                cand = path.with_suffix(f".{i}{path.suffix}")
                if not cand.exists():
                    path = cand
                    break

        size = await _build_docx(article, path)
        return path, size, article.get("title", "")

    try:
        return await asyncio.wait_for(_do(), timeout=TOTAL_TIMEOUT_SECONDS)
    except asyncio.TimeoutError as e:
        raise LinkFetchError(
            f"总耗时超过 {TOTAL_TIMEOUT_SECONDS}s,放弃: {url}"
        ) from e
